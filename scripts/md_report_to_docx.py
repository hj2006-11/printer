"""Convert a daily report markdown file to a Word document.

Usage:
    python scripts/md_report_to_docx.py <input.md> <output.docx>

Supports: headings, paragraphs, bullet lists, numbered lists, images, bold.
Tables are intentionally rendered as plain paragraphs (no Word tables).
"""
import argparse
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_chinese_font(run, font_name="宋体", size=12, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def parse_inline(text):
    """Return list of (text, bold) segments for **bold** markers."""
    segments = []
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            segments.append((part[2:-2], True))
        else:
            segments.append((part, False))
    return segments


def add_runs_to_paragraph(p, text, font_name="宋体", size=12):
    for seg_text, bold in parse_inline(text):
        if seg_text:
            run = p.add_run(seg_text)
            set_chinese_font(run, font_name=font_name, size=size, bold=bold)


def add_heading(doc, text, level):
    style = f"Heading{min(level, 4)}"
    p = doc.add_paragraph(style=style)
    size = {1: 18, 2: 16, 3: 14, 4: 12}.get(level, 12)
    add_runs_to_paragraph(p, text, font_name="黑体", size=size)
    for run in p.runs:
        run.font.bold = True


def add_paragraph(doc, text, first_line_indent=0):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    add_runs_to_paragraph(p, text)


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    add_runs_to_paragraph(p, text)


def add_numbered(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    cleaned = re.sub(r"^\d+\.\s*", "", text)
    add_runs_to_paragraph(p, cleaned)


def add_image(doc, md_dir, alt, src):
    if not os.path.isabs(src):
        src = os.path.normpath(os.path.join(md_dir, src))
    if os.path.exists(src):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(src, width=Inches(5.8))
        except Exception as exc:
            add_paragraph(doc, f"[图片加载失败: {alt} ({src}): {exc}]")
    else:
        add_paragraph(doc, f"[图片未找到: {alt} ({src})]")


def convert(md_path, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    md_dir = os.path.dirname(os.path.abspath(md_path))
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # Heading
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # Image
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if m:
            add_image(doc, md_dir, m.group(1), m.group(2))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            add_numbered(doc, m.group(2))
            i += 1
            continue

        # Bullet list
        if stripped.startswith("-") or stripped.startswith("*"):
            text = re.sub(r"^[-*]\s+", "", stripped)
            add_bullet(doc, text)
            i += 1
            continue

        # Indented bullet (sub-item under a list, e.g. "  - 输入：...")
        m = re.match(r"^(\s{2,})([-*])\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)) // 2, 2)
            add_bullet(doc, m.group(3), level=level)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        # Regular paragraph
        add_paragraph(doc, stripped)
        i += 1

    sections = doc.sections[0]
    sections.top_margin = Cm(2.54)
    sections.bottom_margin = Cm(2.54)
    sections.left_margin = Cm(3.17)
    sections.right_margin = Cm(3.17)

    doc.save(out_path)
    print(f"Saved docx: {out_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert markdown report to docx")
    parser.add_argument("input_md", help="Input markdown file")
    parser.add_argument("output_docx", help="Output Word document")
    args = parser.parse_args()
    convert(args.input_md, args.output_docx)
