"""Enhanced markdown -> docx converter supporting tables and code blocks.

Usage:
    python scripts/md_to_docx_v2.py <input.md> <output.docx>

Supports: headings, paragraphs, bullets, numbered lists, bold, code blocks,
tables, blockquotes, images.
"""
import argparse
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name="宋体", size=12, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def parse_inline(text):
    segments = []
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            segments.append((part[2:-2], True))
        else:
            segments.append((part, False))
    return segments


def add_runs_to_paragraph(p, text, font_name="宋体", size=12):
    for seg_text, bold in parse_inline(text):
        if seg_text:
            run = p.add_run(seg_text)
            set_chinese_font(run, font_name=font_name, size=size, bold=bold)


def add_heading(doc, text, level):
    style = f"Heading{min(level, 4)}"
    p = doc.add_paragraph(style=style)
    size = {1: 18, 2: 16, 3: 14, 4: 12}.get(level, 12)
    add_runs_to_paragraph(p, text, font_name="黑体", size=size)
    for run in p.runs:
        run.font.bold = True


def add_paragraph(doc, text, first_line_indent=0):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    add_runs_to_paragraph(p, text)


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    add_runs_to_paragraph(p, text)


def add_numbered(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    cleaned = re.sub(r"^\d+\.\s*", "", text)
    add_runs_to_paragraph(p, cleaned)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    for seg_text, _ in parse_inline(code):
        if seg_text:
            run = p.add_run(seg_text)
            set_chinese_font(run, font_name="Consolas", size=9)
            run.font.name = "Consolas"


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = ""
            text = row[j] if j < len(row) else ""
            add_runs_to_paragraph(cell.paragraphs[0], text, size=10.5)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    add_runs_to_paragraph(p, text)
    for run in p.runs:
        run.font.italic = True


def add_image(doc, md_dir, alt, src):
    if not os.path.isabs(src):
        src = os.path.normpath(os.path.join(md_dir, src))
    if os.path.exists(src):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(src, width=Inches(5.8))
        except Exception as exc:
            add_paragraph(doc, f"[图片加载失败: {alt} ({src}): {exc}]")
    else:
        add_paragraph(doc, f"[图片未找到: {alt} ({src})]")


def convert(md_path, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    md_dir = os.path.dirname(os.path.abspath(md_path))
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # Code block fence
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(stripped)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # Image
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if m:
            add_image(doc, md_dir, m.group(1), m.group(2))
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1].strip()):
            rows = []
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows.append(header)
            i += 2  # skip header separator
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # Blockquote
        if stripped.startswith(">"):
            add_blockquote(doc, re.sub(r"^>\s?", "", stripped))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            add_numbered(doc, m.group(2))
            i += 1
            continue

        # Bullet list
        if stripped.startswith("-") or stripped.startswith("*"):
            text = re.sub(r"^[-*]\s+", "", stripped)
            add_bullet(doc, text)
            i += 1
            continue

        # Indented bullet
        m = re.match(r"^(\s{2,})([-*])\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)) // 2, 2)
            add_bullet(doc, m.group(3), level=level)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        # Regular paragraph
        add_paragraph(doc, stripped)
        i += 1

    sections = doc.sections[0]
    sections.top_margin = Cm(2.54)
    sections.bottom_margin = Cm(2.54)
    sections.left_margin = Cm(3.17)
    sections.right_margin = Cm(3.17)

    doc.save(out_path)
    print(f"Saved docx: {out_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert markdown to docx")
    parser.add_argument("input_md", help="Input markdown file")
    parser.add_argument("output_docx", help="Output Word document")
    args = parser.parse_args()
    convert(args.input_md, args.output_docx)
